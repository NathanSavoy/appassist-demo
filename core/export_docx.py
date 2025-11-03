from __future__ import annotations
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
from utils.io import write_text


def _set_base_styles(doc: Document):
    styles = doc.styles

    name = styles.add_style("Name", WD_STYLE_TYPE.PARAGRAPH)
    name.base_style = styles["Normal"]
    name.font.name = "Palatino"
    name.font.size = Pt(39)
    name.font.bold = False
    name.font.color.rgb = RGBColor(0x24, 0x2c, 0x38)
    name.paragraph_format.space_after = Pt(0)
    name.paragraph_format.space_before = Pt(0)

    section = styles.add_style("Section Title", WD_STYLE_TYPE.PARAGRAPH)
    section.base_style = styles["Normal"]
    section.font.name = "Palatino"
    section.font.size = Pt(15)
    section.font.color.rgb = RGBColor(0x48, 0x59, 0x70)
    section.paragraph_format.space_before = Pt(8)
    section.paragraph_format.space_after = Pt(2)

    job = styles.add_style("Job Title", WD_STYLE_TYPE.PARAGRAPH)
    job.base_style = styles["Normal"]
    job.font.name = "Palatino"
    job.font.size = Pt(13)
    job.font.italic = True
    job.font.color.rgb = RGBColor(0x26, 0x26, 0x26)  # almost black
    job.paragraph_format.space_after = Pt(0)

    description = styles.add_style("Description", WD_STYLE_TYPE.PARAGRAPH)
    description.base_style = styles["List Bullet"]
    description.font.name = "Palatino"
    description.font.size = Pt(12)
    description.font.color.rgb = RGBColor(0x26, 0x26, 0x26)  # almost black
    description.paragraph_format.space_after = Pt(5)

    default = styles.add_style("Default", WD_STYLE_TYPE.PARAGRAPH)
    default.base_style = styles["Normal"]
    default.font.name = "Palatino"
    default.font.size = Pt(12)
    default.font.color.rgb = RGBColor(0x26, 0x26, 0x26)  # almost black
    default.paragraph_format.space_after = Pt(3)
    default.paragraph_format.space_before = Pt(0)




def render_resume_docx(profile: Dict[str, Any], resume_json: Dict[str, Any], out_path: Path) -> Path:
    doc = Document()
    _set_base_styles(doc)

    # --- Set page margins ---
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    # Name / Contact (plain paragraphs; avoid headers/footers)
    p = doc.add_paragraph(style = "Name")
    run = p.add_run(profile.get("full_name", ""))
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    contact_bits = [profile.get("contact", {}).get("email", ""), profile.get("contact", {}).get("phone", ""), profile.get("contact", {}).get("city", "")]
    links = profile.get("contact", {}).get("links", [])
    for l in links:
        contact_bits.append(f"{l.get('url')}")
    c = doc.add_paragraph(style = "Default")
    c.add_run(" • ".join([b for b in contact_bits if b]))
    c.bold = True
    c.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Headline
    headline = resume_json.get("headline")
    if headline:
        hdln = doc.add_paragraph(style = "Section Title")
        runner = hdln.add_run(headline)
        runner.bold = True
        runner.italic = True

    # Experience (only the bullets returned)
    for section in resume_json.get("sections", []):
        title = section.get("title", "Experience")
        doc.add_paragraph(title.upper() + "  –––––––––––––––––––––––––––––––––––––––––––––––", style="Section Title")
        for item in section.get("items", []):
            header = doc.add_paragraph(style = "Job Title")
            header.add_run(item.get("role", "")).bold = True
            emp = item.get("employer", "")
            loc = item.get("location", "")
            dates = item.get("dates", "")
            tail_bits = [b for b in [emp, loc, dates] if b]
            if tail_bits:
                header.add_run(f" — {' | '.join(tail_bits)}")
            # bullets
            for txt in item.get("bullets", []):
                doc.add_paragraph(txt, style="Description")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def render_cover_letter_docx(profile: Dict[str, Any], cl_json: Dict[str, Any], out_path: Path) -> Path:
    doc = Document()
    _set_base_styles(doc)

    # Optional contact block
    top = doc.add_paragraph(profile.get("full_name", ""), style = "Default")
    contact = profile.get("contact", {})
    doc.add_paragraph(
        "\n".join([x for x in [contact.get("email"), contact.get("phone"), contact.get("city")] if x]),
        style = "Default"
    )
    doc.add_paragraph("")

    # Greeting
    greet = cl_json.get("greeting", "Hiring Team")
    doc.add_paragraph(f"Dear {greet}", style = "Default")
    doc.add_paragraph("")


    for para in cl_json.get("body_paragraphs", []):
        doc.add_paragraph(para, style = "Default")

    doc.add_paragraph("")
    doc.add_paragraph("")

    closing = cl_json.get("closing", "Sincerely,")
    #closing = "Sincerely,"
    doc.add_paragraph(closing, style = "Default")
    doc.add_paragraph(profile.get("full_name", ""), style = "Default")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def write_txt_mirrors(resume_text: str, cl_text: str, out_resume: Path, out_cl: Path):
    write_text(out_resume.with_suffix('.txt'), resume_text)
    write_text(out_cl.with_suffix('.txt'), cl_text)
